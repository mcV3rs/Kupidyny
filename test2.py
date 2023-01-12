# Create/Open a document
from docx import Document
from docx.shared import Inches

doc = Document()

# add image and save doc
doc.add_heading("Image", 0)
doc.add_picture("C:\\Users\\mcVer\\Desktop\\SilesianPhoenix\\Kupidyny\\project\\instance\\1.png", width=Inches(5),
                height=Inches(3))
doc.save("doc_images.docx")
